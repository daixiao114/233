from sympy import isprime, primerange

# 选择素数 p 和原根 g
p = 101
g = 2

# A 和 B 分别选择私有密钥
a = 5
b = 7

# 计算公有值
A = pow(g, a, p)
B = pow(g, b, p)



# 生成共享密钥
K_A = pow(B, a, p)
K_B = pow(A, b, p)


from docx import Document

# 创建一个新的Word文档
doc = Document()

# 添加标题
doc.add_heading('D-H Key Exchange Experiment', level=1)

# 添加计算过程
doc.add_paragraph(f"Prime number p: {p}")
doc.add_paragraph(f"Primitive root g: {g}")
doc.add_paragraph(f"A's private key: {a}")
doc.add_paragraph(f"B's private key: {b}")
doc.add_paragraph(f"A's public value: {A}")
doc.add_paragraph(f"B's public value: {B}")
doc.add_paragraph(f"A's computed shared secret: {K_A}")
doc.add_paragraph(f"B's computed shared secret: {K_B}")

# 保存文档
doc.save('dh_key_exchange_experiment.docx')
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas

# 创建一个新的PDF文档
c = canvas.Canvas("dh_key_exchange_experiment.pdf", pagesize=letter)
width, height = letter

# 添加文本
text = c.beginText(40, height - 40)
text.setFont("Helvetica", 12)
text.textLine("D-H Key Exchange Experiment")
text.textLine("")
text.textLines(f"Prime number p: {p}\n"
               f"Primitive root g: {g}\n"
               f"A's private key: {a}\n"
               f"B's private key: {b}\n"
               f"A's public value: {A}\n"
               f"B's public value: {B}\n"
               f"A's computed shared secret: {K_A}\n"
               f"B's computed shared secret: {K_B}")
c.drawText(text)

# 保存文档
c.showPage()
c.save()